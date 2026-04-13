"""
Build Project3_CloseTheGap_InterviewGuide.docx using python-docx.
Run: python3 build_interview_guide.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x6B)   # deep navy
TEAL   = RGBColor(0x00, 0x7A, 0x87)   # HACR teal accent
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xE8, 0xF4, 0xF8)   # very light blue
DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black body text
AMBER  = RGBColor(0xE6, 0x8A, 0x00)   # highlight accent
MID    = RGBColor(0xF2, 0xF8, 0xFB)   # table alt row

OUTPUT = "/Users/yamini/Documents/Housing-Project/Project3_CloseTheGap_InterviewGuide.docx"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_para_border_bottom(para, color="007A87", size=12):
    """Draw a coloured bottom border under a paragraph (used for section dividers)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def shade_para(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def styled_run(para, text, bold=False, italic=False,
               color=None, size=None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.font.name = font
    return run

def section_heading(doc, text, level=1):
    """Navy-coloured section heading with teal bottom border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)
    run  = para.add_run(text)
    run.bold            = True
    run.font.color.rgb  = NAVY
    run.font.name       = "Calibri"
    run.font.size       = Pt(14) if level == 1 else Pt(12)
    add_para_border_bottom(para, color="007A87", size=8 if level == 1 else 4)
    return para

def sub_heading(doc, text):
    """Teal sub-heading, no border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(text)
    run.bold           = True
    run.font.color.rgb = TEAL
    run.font.name      = "Calibri"
    run.font.size      = Pt(11)
    return para

def bullet(doc, text, indent=0.35):
    """Clean bullet point."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent   = Inches(indent)
    para.paragraph_format.space_before  = Pt(1)
    para.paragraph_format.space_after   = Pt(1)
    run = para.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(10.5)
    run.font.color.rgb = DARK
    return para

def body(doc, text, italic=False, color=None, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.italic         = italic
    run.font.name      = "Calibri"
    run.font.size      = Pt(10.5)
    run.font.color.rgb = color or DARK
    return para

# ── Build document ───────────────────────────────────────────────────────────
doc = Document()

# Page margins — 1 inch all around
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ════════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(6)
title_para.paragraph_format.space_after  = Pt(2)
shade_para(title_para, "1B3A6B")

t1 = title_para.add_run('"Closing the Gap"  —  Interview Prep Guide')
t1.bold           = True
t1.font.size      = Pt(18)
t1.font.color.rgb = WHITE
t1.font.name      = "Calibri"

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_after = Pt(10)
shade_para(sub_para, "007A87")
t2 = sub_para.add_run(
    "Wake County Housing Supply-Demand Dashboard  |  Built with Real HACR Data"
)
t2.bold           = True
t2.font.size      = Pt(11)
t2.font.color.rgb = WHITE
t2.font.name      = "Calibri"

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — What Is This Project?
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "SECTION 1 — What Is This Project?  (Plain English)")

body(doc,
     "This is an interactive data dashboard built in Python using a tool called Streamlit "
     "(similar to Power BI or Tableau). It uses four real Excel files downloaded from "
     "wakehousingdata.org — the same data Wake County HACR publishes publicly. The dashboard "
     "answers one core question: \u201cHow many affordable homes does Wake County need, how long "
     "will it take to close the gap, and what will it cost?\u201d")

sub_heading(doc, "The Dashboard Has 5 Tabs")

tabs = [
    ("Tab 1 — The Gap:",          "Shows how many affordable units are missing by income level."),
    ("Tab 2 — Affordability Squeeze:",  "Shows how rent, income, and housing costs have drifted apart over time."),
    ("Tab 3 — Supply Pressure:",  "Shows how tight the housing market is (vacancy rates vs. healthy benchmarks)."),
    ("Tab 4 — Scenario Planner:", "Lets you model how fast the gap could close under different policy choices."),
    ("Tab 5 — CAPER Export:",     "Exports data in the format used in federal HUD reports submitted by Wake County."),
]
for label, desc in tabs:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent  = Inches(0.35)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    r1 = para.add_run(label + "  ")
    r1.bold            = True
    r1.font.color.rgb  = TEAL
    r1.font.name       = "Calibri"
    r1.font.size       = Pt(10.5)
    r2 = para.add_run(desc)
    r2.font.name       = "Calibri"
    r2.font.size       = Pt(10.5)
    r2.font.color.rgb  = DARK

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Key Findings Table
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "SECTION 2 — The Real Data: Key Findings")

body(doc,
     "Every number below comes directly from the four wakehousingdata.org Excel files. "
     "These are real Wake County figures — not estimates or synthetic data.", space_after=6)

findings = [
    # (Finding, The Number, What It Means)
    (
        "Affordable unit shortage for households earning under $35K",
        "26,037 units missing",
        "Wake County has only 33 affordable units for every 100 renter households at this income "
        "level — the rest have nowhere affordable to go."
    ),
    (
        "Years to close the gap at current HACR production pace",
        "39 years  (closes in 2065)",
        "HACR produced 660 affordable units in FY2025. At that rate, the shortage won't be "
        "resolved until 2065."
    ),
    (
        "Rent spike in a single year (2021)",
        "Income needed jumped $8,920 in one year",
        "Largest single-year affordability shock in the dataset — rents jumped $223/month while "
        "incomes barely moved."
    ),
    (
        "Collapse in entry-level home sales",
        "Under-$250K homes: 28.6% of sales (2018)  →  3.1% (2022)  =  89% drop",
        "First-time buyers were essentially priced out of the Wake County market in just 4 years."
    ),
    (
        "New income bands pushed into housing crisis",
        "$35K–$50K cost burden: 89%  |  $50K–$75K: 56%",
        "The affordability crisis is no longer just a low-income problem — it has moved up the "
        "income ladder."
    ),
    (
        "Vacancy rate in Wake County (2024)",
        "3.72%  (healthy benchmark = 5%)",
        "When vacancy is this low, landlords can raise rents freely — this directly drives cost "
        "burden higher."
    ),
    (
        "Total investment needed to close the <$35K gap",
        "$3.8 billion — same regardless of pace",
        "Whether the gap closes in 8 years or 40 years, the total dollars needed are the same. "
        "Only the annual commitment (and urgency) changes."
    ),
]

# Table: 3 columns — Finding | The Number | What It Means
col_w = [Inches(2.1), Inches(1.8), Inches(2.9)]
tbl   = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"

# Header row
hdr_cells = tbl.rows[0].cells
headers   = ["Finding", "The Number", "What It Means"]
hdr_colors = ["1B3A6B", "1B3A6B", "1B3A6B"]
for i, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
    set_cell_bg(cell, hdr_colors[i])
    set_cell_borders(cell, "FFFFFF")
    cell.width = col_w[i]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(hdr)
    r.bold            = True
    r.font.color.rgb  = WHITE
    r.font.name       = "Calibri"
    r.font.size       = Pt(10)

# Data rows
for idx, (finding, number, meaning) in enumerate(findings):
    row       = tbl.add_row()
    row_cells = row.cells
    bg = "E8F4F8" if idx % 2 == 0 else "F7FBFD"
    for cell in row_cells:
        set_cell_bg(cell, bg)
        set_cell_borders(cell, "B8D8E8")

    # Col 0 — Finding
    row_cells[0].width = col_w[0]
    p0 = row_cells[0].paragraphs[0]
    p0.paragraph_format.space_before = Pt(3)
    p0.paragraph_format.space_after  = Pt(3)
    r0 = p0.add_run(finding)
    r0.bold           = True
    r0.font.name      = "Calibri"
    r0.font.size      = Pt(9.5)
    r0.font.color.rgb = NAVY

    # Col 1 — The Number
    row_cells[1].width = col_w[1]
    p1 = row_cells[1].paragraphs[0]
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after  = Pt(3)
    r1 = p1.add_run(number)
    r1.bold           = True
    r1.font.name      = "Calibri"
    r1.font.size      = Pt(9.5)
    r1.font.color.rgb = AMBER

    # Col 2 — Meaning
    row_cells[2].width = col_w[2]
    p2 = row_cells[2].paragraphs[0]
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after  = Pt(3)
    r2 = p2.add_run(meaning)
    r2.font.name      = "Calibri"
    r2.font.size      = Pt(9.5)
    r2.font.color.rgb = DARK

doc.add_paragraph()  # spacer

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — How the Dashboard Works
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "SECTION 3 — How the Dashboard Works  (Technical, Simple Terms)")

sub_heading(doc, "Data Sources")
bullet(doc, "Four Excel files from wakehousingdata.org covering: Rental Affordability, Homeownership, Demographics, and Housing Supply.")
bullet(doc, "All numbers match Wake County's 2025 Annual Report benchmarks: 7,086 units produced to date, 660 units/year in FY2025, $18M annual investment, 2,500-unit production goal.")

sub_heading(doc, "The Gap Formula")
bullet(doc, "Gap = Affordable Units Available  MINUS  Renter Households at That Income Level.")
bullet(doc, "A negative number means a shortage. Example: at the <$35K band there are 26,037 fewer affordable units than renter households who need them.")
bullet(doc, "Source: HUD's Worst Case Housing Needs methodology (affordable = unit where rent + utilities is 30% or less of gross income).")

sub_heading(doc, "Years-to-Close Formula")
bullet(doc, "Years to Close  =  Gap  ÷  (Annual Production  ×  % of units targeting that income band)")
bullet(doc, "At current pace:  26,037  ÷  (660 × 1.0)  =  ~39.5 years")
bullet(doc, "At 5× scale-up (3,300 units/year):  26,037  ÷  3,300  =  ~8 years")

sub_heading(doc, "Scenario Planner (Tab 4)")
bullet(doc, "4 named policy scenarios: Current Pace, 2× Scale-Up, 5× Scale-Up, and Targeted Deep Subsidy.")
bullet(doc, "Key insight: The total cost ($3.8B) is the same no matter the pace — doubling production halves the timeline but does not change total spend. Only the annual commitment (and the urgency) differs.")
bullet(doc, "A sensitivity table shows years-to-close for every income band across 7 different production rate options — giving planners an at-a-glance policy menu.")

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — How to Talk About This in the Interview
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "SECTION 4 — How to Talk About This in the Interview")

sub_heading(doc, 'Opening Pitch  (when asked "tell me about a project you\'ve done")')
quote_para = doc.add_paragraph()
quote_para.paragraph_format.left_indent  = Inches(0.4)
quote_para.paragraph_format.right_indent = Inches(0.4)
quote_para.paragraph_format.space_before = Pt(4)
quote_para.paragraph_format.space_after  = Pt(8)
shade_para(quote_para, "E8F4F8")
qr = quote_para.add_run(
    "\u201cI built an interactive dashboard using real Wake County HACR data from wakehousingdata.org. "
    "It maps the full affordable housing supply gap by income band, models policy scenarios showing "
    "how long it would take to close the gap under different investment levels, and exports data in "
    "a format compatible with HUD CAPER reporting. The headline finding is that Wake County is "
    "26,037 units short for households earning under $35,000 \u2014 and at the current production "
    "pace, the gap won\u2019t close until 2065.\u201d"
)
qr.italic          = True
qr.font.name       = "Calibri"
qr.font.size       = Pt(10.5)
qr.font.color.rgb  = DARK

sub_heading(doc, "If Asked About Methodology")
bullet(doc, "Used HUD's cost burden definition: a household is cost-burdened if they spend more than 30% of gross income on housing (24 CFR §5.609).")
bullet(doc, "AMI = Area Median Income — HUD sets this annually for every metro area; it is the standard benchmark for affordable housing eligibility.")
bullet(doc, "Gap calculation follows HUD's Worst Case Housing Needs methodology.")
bullet(doc, "Production data sourced directly from HACR's FY2025 Annual Report.")

sub_heading(doc, 'If Asked "What Would You Improve or Do Next?"')
bullet(doc, "Add census tract-level mapping to show which Wake County neighbourhoods have the worst shortage (geographic equity lens).")
bullet(doc, "Layer in race and ethnicity data — the homeownership data shows Black households own at 45.5% vs 71.4% for White households in Wake County; that disparity deserves a dedicated view.")
bullet(doc, "Connect to live HMIS data to show how homelessness outcomes correlate with the supply gap in real time.")
bullet(doc, 'Add a "what if" builder where HACR staff can input a new proposed development and instantly see its impact on the gap.')

sub_heading(doc, "Key Terms to Use Naturally in the Interview")

terms = [
    ('"AMI band"',   'Income tier expressed as a % of Area Median Income (e.g., 30% AMI, 50% AMI, 80% AMI).'),
    ('"Cost burden"', 'Spending more than 30% of income on housing. "Severely cost burdened" means more than 50%.'),
    ('"CAPER"',      'Consolidated Annual Performance and Evaluation Report — the federal report Wake County submits to HUD each year to account for grant spending.'),
    ('"Pipeline units"', 'Affordable units that are approved but not yet built. Wake County had 1,532 in the pipeline as of FY2025.'),
    ('"LIHTC"',      'Low Income Housing Tax Credit — the primary federal financing mechanism for affordable housing development.'),
]
for term, defn in terms:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent  = Inches(0.35)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    r1 = para.add_run(term + "  \u2014  ")
    r1.bold            = True
    r1.font.color.rgb  = TEAL
    r1.font.name       = "Calibri"
    r1.font.size       = Pt(10.5)
    r2 = para.add_run(defn)
    r2.font.name       = "Calibri"
    r2.font.size       = Pt(10.5)
    r2.font.color.rgb  = DARK

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — The 4 Excel Files
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "SECTION 5 — The 4 Excel Files: What's in Each One")

files_data = [
    (
        "1.  Rental Affordability",
        [
            "Median rent trends (2015–2024): Wake County rent rose from $1,104 to $1,648 — a 49% increase in 9 years.",
            "Rental deficit/surplus by income band: 26,037-unit shortage at the <$35K income level.",
            "Renter cost burden by income: 96% of renters earning under $20K are cost-burdened in 2024.",
            "Income vs. rent required: Renters needed $44,160 income to afford median rent in 2015; by 2024 that rose to $65,920.",
        ]
    ),
    (
        "2.  Homeownership",
        [
            "Median home value: Rose from $229,200 (2014) to $461,300 (2024) — doubled in 10 years.",
            "Homeownership rate by race: Asian 72.0%  |  White 71.4%  |  Hispanic 49.9%  |  Black 45.5%.",
            "Entry-level homes (<$250K) collapsed from 28.6% of all sales in 2018 to just 3.1% in 2022.",
            "Owner cost burden: 80.7% of homeowners earning under $20K are severely cost-burdened in 2024.",
        ]
    ),
    (
        "3.  Demographics",
        [
            "Wake County population: Grew from 1.02M (2015) to 1.26M (2025) — 22.8% growth in a decade.",
            "Households earning under $35K fell by 23,176 between 2014 and 2024 — not because they got richer, but because they were priced out.",
            "Income distribution shifted sharply: households earning $150K+ grew by 100,784; those under $75K shrank.",
            "Racial composition of households: White 283,141  |  Black 88,430  |  Asian 33,593  |  Hispanic/Latino 38,091.",
        ]
    ),
    (
        "4.  Housing Supply",
        [
            "HACR has produced 7,086 affordable units to date (cumulative).",
            "FY2025 production: 660 units delivered; goal was 2,500 units (achieved 2,065 including pipeline).",
            "$18 million annual investment from Wake County into affordable housing.",
            "Vacancy rate: 3.72% in 2024 — below the 5% benchmark considered healthy for a balanced market.",
        ]
    ),
]

for file_title, bullets_list in files_data:
    sub_heading(doc, file_title)
    for b in bullets_list:
        bullet(doc, b)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — What Makes This Project Stand Out
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "SECTION 6 — What Makes This Project Stand Out")

standout = [
    "Uses real public data — not made-up numbers. Every figure matches HACR's own published reports and website.",
    "Built in Python (Streamlit + Plotly) — demonstrates technical skill without needing a paid tool licence like Tableau or Power BI.",
    "CAPER export feature — directly relevant to the job; shows awareness of federal reporting requirements that HACR staff deal with daily.",
    "Scenario modelling — shows policy thinking, not just data display. The interviewer can change production rates and watch the timeline shift in real time.",
    "The \"39 years to close the gap\" finding is a concrete, memorable statistic that connects data to urgency in a way any audience understands.",
    "Fully interactive — the dashboard can be launched live on a laptop during the interview, showing rather than just telling.",
]
for s in standout:
    bullet(doc, s)

# ════════════════════════════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_para(footer_para, "1B3A6B")
fr = footer_para.add_run(
    "All data sourced from wakehousingdata.org (Wake County HACR Public Dashboard, downloaded April 2026).  "
    "Dashboard built April 2026 in Python 3.9 using Streamlit, Plotly, pandas, and openpyxl."
)
fr.font.size      = Pt(8)
fr.font.color.rgb = WHITE
fr.font.name      = "Calibri"
fr.italic         = True

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"✓ Saved: {OUTPUT}")
