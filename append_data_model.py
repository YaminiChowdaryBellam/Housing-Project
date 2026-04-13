"""
Append a Data Model section to Project3_CloseTheGap_InterviewGuide.docx.
Run: python3 append_data_model.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x1B, 0x3A, 0x6B)
TEAL  = RGBColor(0x00, 0x7A, 0x87)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
AMBER = RGBColor(0xE6, 0x8A, 0x00)
GREEN = RGBColor(0x1A, 0x7A, 0x3C)
GREY  = RGBColor(0x55, 0x55, 0x55)

PATH = "/Users/yamini/Documents/Housing-Project/Project3_CloseTheGap_InterviewGuide.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_bottom_border(para, color="007A87", size=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def shade_para(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def section_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    add_bottom_border(para)
    return para

def sub_heading(doc, text, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = color or TEAL
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return para

def body(doc, text, italic=False, color=None, space_after=4, indent=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    if indent:
        para.paragraph_format.left_indent = Inches(indent)
    run = para.add_run(text)
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or DARK
    return para

def bullet(doc, text, indent=0.35):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    return para

def two_col_bullet(doc, label, desc, label_color=None):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.35)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    r1 = para.add_run(label + "  ")
    r1.bold = True
    r1.font.color.rgb = label_color or TEAL
    r1.font.name = "Calibri"
    r1.font.size = Pt(10.5)
    r2 = para.add_run(desc)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK

def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(__import__("docx.oxml.ns", fromlist=["qn"])
                  .__class__)  # placeholder — use OxmlElement
    from docx.oxml import OxmlElement as OE
    br = OE("w:br")
    br.set(qn("w:type"), "page")
    para._p.append(br)

# ── load existing doc ─────────────────────────────────────────────────────────
doc = Document(PATH)

# Page break before new section
pb_para = doc.add_paragraph()
from docx.oxml import OxmlElement as OE
br = OE("w:br")
br.set(qn("w:type"), "page")
pb_para._p.append(br)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DATA MODEL
# ════════════════════════════════════════════════════════════════════════════

# Section title banner
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(0)
title_para.paragraph_format.space_after = Pt(2)
shade_para(title_para, "1B3A6B")
tr = title_para.add_run("SECTION 7 — Data Model & Table Relationships")
tr.bold = True
tr.font.size = Pt(15)
tr.font.color.rgb = WHITE
tr.font.name = "Calibri"

sub_banner = doc.add_paragraph()
sub_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_banner.paragraph_format.space_after = Pt(10)
shade_para(sub_banner, "007A87")
sr = sub_banner.add_run(
    "How the 4 wakehousingdata.org Excel files connect to each other"
)
sr.bold = True
sr.font.size = Pt(10)
sr.font.color.rgb = WHITE
sr.font.name = "Calibri"

# ── 7.1 Overview ─────────────────────────────────────────────────────────────
section_heading(doc, "7.1  Overview — The 4 Source Files")

body(doc,
     "All data in this dashboard comes from four Excel files published by Wake County HACR at "
     "wakehousingdata.org. Think of each file as a subject area. They are not formally joined "
     "like a SQL database, but they share common dimensions — year, income band, and geography "
     "— that allow them to be analysed together.",
     space_after=6)

# Overview table: File | Subject | Key Tables Inside | Grain
overview_rows = [
    ("Demographics",        "Who lives in Wake County",
     "Population, household income, race/ethnicity, age, education, household size",
     "Year  |  Income band  |  Race  |  Age group"),
    ("Homeownership",       "Can residents buy a home?",
     "Home values, homeownership rates, owner cost burden, entry-level sales, vacancy",
     "Year  |  Income band  |  Race"),
    ("Rental Affordability","Can residents rent a home?",
     "Median rent, rental deficit, affordable units per 100, renter cost burden, income vs. rent",
     "Year  |  Income band  |  Education"),
    ("Housing Supply",      "Is enough housing being built?",
     "Building permits, vacancy rate, stock by tenure, demand vs. supply by income",
     "Year  |  Income band  |  Building type"),
]

tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
hdrs = ["File / Subject Area", "What It Covers", "Key Tables Inside", "Shared Dimensions"]
hdr_w = [Inches(1.5), Inches(1.5), Inches(2.5), Inches(1.8)]
for i, (cell, h) in enumerate(zip(tbl.rows[0].cells, hdrs)):
    set_cell_bg(cell, "1B3A6B")
    set_cell_borders(cell, "FFFFFF")
    cell.width = hdr_w[i]
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.name = "Calibri"
    r.font.size = Pt(9)

for idx, (file, covers, tables, dims) in enumerate(overview_rows):
    row = tbl.add_row()
    bg = "E8F4F8" if idx % 2 == 0 else "F7FBFD"
    vals = [file, covers, tables, dims]
    colors = [NAVY, DARK, GREY, GREEN]
    bolds  = [True, False, False, False]
    for ci, (cell, val) in enumerate(zip(row.cells, vals)):
        set_cell_bg(cell, bg)
        set_cell_borders(cell, "B8D8E8")
        cell.width = hdr_w[ci]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(val)
        r.bold = bolds[ci]
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = colors[ci]

doc.add_paragraph()

# ── 7.2 Shared Dimensions ────────────────────────────────────────────────────
section_heading(doc, "7.2  The 3 Shared Dimensions (How Tables Talk to Each Other)")

body(doc,
     "Even though these are separate Excel files, three common dimensions act as the glue "
     "that links them. When you compare numbers across files, you always join on one of these:",
     space_after=6)

dims_tbl = doc.add_table(rows=1, cols=3)
dims_tbl.style = "Table Grid"
for cell, h in zip(dims_tbl.rows[0].cells, ["Dimension", "Values in the Data", "Which Files Use It"]):
    set_cell_bg(cell, "007A87")
    set_cell_borders(cell, "FFFFFF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)

dim_rows = [
    ("Year",
     "2013–2025 (varies by table; most run 2014–2024)",
     "All 4 files — time-series data in every file uses Year as the row key"),
    ("Income Band",
     "<$20K  |  $20K–$35K  |  $35K–$50K  |  $50K–$75K  |  $75K–$100K  |  $100K–$150K  |  $150K+",
     "Demographics (income distribution, tenure) · Rental (deficit, cost burden) · Homeownership (cost burden, ownership rate) · Housing Supply (demand vs. supply)"),
    ("Race / Ethnicity",
     "White  |  Black  |  Asian  |  Hispanic/Latino  |  AI/AN  |  NH/PI  |  Other",
     "Demographics (households by race) · Homeownership (ownership rate by race)"),
]

for idx, (dim, vals, files) in enumerate(dim_rows):
    row = dims_tbl.add_row()
    bg = "E8F4F8" if idx % 2 == 0 else "F7FBFD"
    data = [dim, vals, files]
    fcolors = [NAVY, DARK, GREY]
    fbolds  = [True, False, False]
    widths  = [Inches(1.3), Inches(3.2), Inches(2.8)]
    for ci, (cell, val) in enumerate(zip(row.cells, data)):
        set_cell_bg(cell, bg)
        set_cell_borders(cell, "B8D8E8")
        cell.width = widths[ci]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(val)
        r.bold = fbolds[ci]
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = fcolors[ci]

doc.add_paragraph()

# ── 7.3 Table-level map ───────────────────────────────────────────────────────
section_heading(doc, "7.3  Every Table and What It Measures")

body(doc,
     "Below is every individual data table in the four files, what it measures, "
     "its grain (one row = one what?), and which other tables it connects to.",
     space_after=6)

all_tables = [
    # (File, Table Name, Measures, Grain, Connects To)
    # ── DEMOGRAPHICS ──
    ("Demographics", "Population Over Time",
     "Total Wake County population by year",
     "1 row = 1 year",
     "Building Permits (supply keeps pace?); Net HH Change (growth context)"),
    ("Demographics", "% Change in Population",
     "Population growth rate — Wake vs. NC",
     "1 row = 1 year × 2 geographies",
     "Vacancy Rate (demand pressure); Median Rent trend"),
    ("Demographics", "Household Income Distribution",
     "Share of households in each income band — 2014 vs 2024",
     "1 row = 1 income band",
     "Rental Deficit by Income; Owner Cost Burden by Income"),
    ("Demographics", "Net Change in HH by Income Band",
     "How many more/fewer households at each income level 2014→2024",
     "1 row = 1 income band",
     "Rental Deficit (disappearing demand vs. priced-out); HH by Income & Tenure"),
    ("Demographics", "Households by Income & Tenure",
     "What % of owners vs. renters fall in each income band",
     "1 row = 1 income band",
     "Renter Cost Burden by Income; Owner Cost Burden by Income"),
    ("Demographics", "Households by Race/Ethnicity",
     "Total household count by racial group",
     "1 row = 1 race group",
     "Homeownership Rate by Race (equity gap)"),
    ("Demographics", "Population by Age",
     "Age group breakdown — 2014, 2019, 2024",
     "1 row = 1 age group × 3 years",
     "Household Size; first-time buyer context"),
    ("Demographics", "Median Household Income",
     "Median income — Wake vs. NC, 2014–2024",
     "1 row = 1 year × 2 geographies",
     "Income Required to Afford Rent (affordability gap)"),
    ("Demographics", "Educational Attainment",
     "Share of population by education level — 2014 vs 2024",
     "1 row = 1 education level",
     "Rent Affordable by Education (Tab 2 chart)"),

    # ── HOMEOWNERSHIP ──
    ("Homeownership", "Homeownership Rate",
     "% of households that own — Wake vs. NC, 2014–2024",
     "1 row = 1 year × 2 geographies",
     "Homeownership Rate by Income; by Race"),
    ("Homeownership", "Median Home Values",
     "Median sale price — Wake vs. NC, 2014–2024",
     "1 row = 1 year × 2 geographies",
     "Household Income (affordability ratio); Entry-Level Sales"),
    ("Homeownership", "Owner Cost Burden by Income",
     "% severely cost burdened & cost burdened owners — by income band, 3 years",
     "1 row = 1 income band × 3 years (2014/2019/2024)",
     "Renter Cost Burden by Income (compare owner vs. renter hardship)"),
    ("Homeownership", "Homes Sold by Price Bracket",
     "Share of sales in each price range — 2018 vs 2022",
     "1 row = 1 price bracket × 2 years",
     "Median Home Value; Household Income Distribution"),
    ("Homeownership", "Homeownership Rate by Income",
     "% who own at each income level — Wake vs. NC",
     "1 row = 1 income band × 2 geographies",
     "Households by Income & Tenure; Rental Deficit (renters who can't buy)"),
    ("Homeownership", "Homeownership Rate by Race",
     "% who own by racial group — Wake vs. NC",
     "1 row = 1 race group × 2 geographies",
     "Households by Race/Ethnicity (equity gap analysis)"),
    ("Homeownership", "Homeowner Vacancy Rate",
     "% of owner homes vacant — Wake vs. NC, 2014–2024",
     "1 row = 1 year × 2 geographies",
     "Rental Vacancy Rate (overall market tightness)"),

    # ── RENTAL AFFORDABILITY ──
    ("Rental", "Median Rent",
     "Median gross rent — Wake vs. NC, 2015–2024",
     "1 row = 1 year × 2 geographies",
     "Median Household Income; Income Required to Afford Rent"),
    ("Rental", "Rental Deficit/Surplus by Income",
     "Gap between affordable units available and renter households needing them",
     "1 row = 1 income band (cumulative thresholds)",
     "Demand vs. Supply by Income (Housing Supply file — same numbers, different view)"),
    ("Rental", "Affordable Units per 100 Renters",
     "How many affordable homes exist for every 100 renters at each income level",
     "1 row = 1 income band",
     "Rental Deficit (derived from same source data)"),
    ("Rental", "Renter Cost Burden by Income",
     "% of renters cost-burdened — by income band — 2014, 2019, 2024",
     "1 row = 1 income band × 3 years",
     "Owner Cost Burden by Income (compare tenures); Net HH Change (bands in crisis)"),
    ("Rental", "Rent Affordable by Education",
     "What rent can each education level afford vs. actual median rent — 2023",
     "1 row = 1 education level",
     "Educational Attainment (Demographics); Median Rent"),
    ("Rental", "Income Required vs. Renter Income",
     "Income needed to afford median rent vs. actual median renter income — 2015–2024",
     "1 row = 1 year",
     "Median Household Income; Median Rent (3-way affordability view)"),

    # ── HOUSING SUPPLY ──
    ("Housing Supply", "Building Permits Issued",
     "Annual single-family and multifamily permits — Wake, 2014–2024",
     "1 row = 1 year × 2 permit types",
     "% Change in Total Homes; HACR production benchmarks"),
    ("Housing Supply", "Vacancy Rate (Rental)",
     "% of rental homes vacant and available — Wake vs. NC, 2014–2024",
     "1 row = 1 year × 2 geographies",
     "Median Rent (low vacancy → rent pressure); Homeowner Vacancy Rate"),
    ("Housing Supply", "Demand vs. Supply by Income",
     "Renter households vs. affordable units available — by income band",
     "1 row = 1 income band",
     "Rental Deficit (identical figures); Renter Cost Burden (explains why burden is high)"),
    ("Housing Supply", "Housing Stock by Tenure",
     "Total occupied, vacant, owner, and renter units — snapshot",
     "1 row = 1 tenure category",
     "Homeownership Rate; Vacancy Rate"),
    ("Housing Supply", "Stock by Building Type",
     "Owner and renter homes by building typology — Wake vs. NC",
     "1 row = 1 building type × tenure × 2 geographies",
     "Permits by type (single-family vs. multifamily trend)"),
]

# Group by file
from collections import OrderedDict
file_groups = OrderedDict()
for row in all_tables:
    file_groups.setdefault(row[0], []).append(row[1:])

file_colors = {
    "Demographics":        ("1B3A6B", "E8EEF8"),
    "Homeownership":       ("007A87", "E8F4F8"),
    "Rental":              ("5B2D8E", "F0EAF8"),
    "Housing Supply":      ("1A7A3C", "EAF4EE"),
}

for file_name, rows in file_groups.items():
    hdr_hex, row_hex = file_colors[file_name]

    # File sub-heading
    fh_para = doc.add_paragraph()
    fh_para.paragraph_format.space_before = Pt(10)
    fh_para.paragraph_format.space_after = Pt(3)
    shade_para(fh_para, hdr_hex)
    fhr = fh_para.add_run(f"  {file_name} File")
    fhr.bold = True
    fhr.font.color.rgb = WHITE
    fhr.font.name = "Calibri"
    fhr.font.size = Pt(11)

    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    col_hdrs = ["Table Name", "What It Measures", "Grain (1 row = ...)", "Connects To"]
    col_w2   = [Inches(1.6), Inches(2.4), Inches(1.6), Inches(1.7)]
    for ci, (cell, h) in enumerate(zip(tbl2.rows[0].cells, col_hdrs)):
        set_cell_bg(cell, hdr_hex)
        set_cell_borders(cell, "FFFFFF")
        cell.width = col_w2[ci]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"
        r.font.size = Pt(8.5)

    for ri, (tname, measures, grain, connects) in enumerate(rows):
        row2 = tbl2.add_row()
        bg2 = row_hex if ri % 2 == 0 else "FAFCFF"
        vals2   = [tname, measures, grain, connects]
        fcolors2 = [NAVY, DARK, GREY, GREEN]
        fbolds2  = [True, False, False, False]
        for ci2, (cell2, val2) in enumerate(zip(row2.cells, vals2)):
            set_cell_bg(cell2, bg2)
            set_cell_borders(cell2, "CCDDEE")
            cell2.width = col_w2[ci2]
            p2 = cell2.paragraphs[0]
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after = Pt(2)
            r2 = p2.add_run(val2)
            r2.bold = fbolds2[ci2]
            r2.font.name = "Calibri"
            r2.font.size = Pt(8.5)
            r2.font.color.rgb = fcolors2[ci2]

    doc.add_paragraph()

# ── 7.4 Key Cross-File Relationships ─────────────────────────────────────────
section_heading(doc, "7.4  Key Cross-File Relationships Explained Simply")

body(doc,
     "These are the six most important connections between tables — the ones the dashboard "
     "actually uses to build its charts and findings.",
     space_after=6)

relationships = [
    (
        "1.  Rental Deficit  ←→  Demand vs. Supply (Housing Supply)",
        "These two tables show the same gap from two angles. "
        "\"Rental Deficit\" (Rental file) shows just the net gap number. "
        "\"Demand vs. Supply\" (Housing Supply file) shows both sides — "
        "how many renter households exist AND how many affordable units exist. "
        "Together they explain why the gap is the size it is.",
        "Same income bands (<$20K through <$150K). Join on income band."
    ),
    (
        "2.  Renter Cost Burden  ←→  Net HH Change (Demographics)",
        "As cost burden rises in a band (Rental file), households at that income "
        "level disappear from Wake County (Demographics file). "
        "The $35K\u2013$50K band is the clearest example: cost burden jumped from 34.5% "
        "to 88.9%, and that band lost 7,177 households.",
        "Same income bands. Join on income band."
    ),
    (
        "3.  Median Rent  ←→  Median Household Income  ←→  Income Required",
        "Three tables tracking the same affordability squeeze from different angles: "
        "actual rent (Rental file), actual renter income (Demographics file), "
        "and the income you would need to afford the rent (Rental file). "
        "When the gap between lines 2 and 3 widens, more people become cost-burdened.",
        "Same years (2015\u20132024). Join on year."
    ),
    (
        "4.  Educational Attainment (Demographics)  ←→  Rent by Education (Rental)",
        "The education split shows what rent each degree level can afford. "
        "Combined with how many Wake County residents hold each degree "
        "(Demographics file), you can see exactly which education groups "
        "cannot afford median rent \u2014 and how large that population is.",
        "Same education levels. Join on education category."
    ),
    (
        "5.  Vacancy Rate (Housing Supply)  ←→  Median Rent (Rental)",
        "Low vacancy means landlords have no pressure to keep rents competitive. "
        "Wake County's vacancy rate fell from 3.58% (2014) to 2.74% (2019), "
        "exactly the period when rent acceleration began. "
        "By 2022 vacancy jumped to 3.77% as multifamily building surged \u2014 "
        "and rent growth slowed.",
        "Same years. Join on year."
    ),
    (
        "6.  Homeownership Rate by Race (Homeownership)  ←→  Households by Race (Demographics)",
        "The race gap in homeownership (Black 45.5% vs White 71.4%) is the rate. "
        "The Demographics file gives the absolute household counts by race. "
        "Multiplying rate \u00d7 count gives the actual number of households "
        "locked out of ownership \u2014 the equity gap in real numbers.",
        "Same race categories. Join on race/ethnicity."
    ),
]

for title, explanation, join_key in relationships:
    rel_para = doc.add_paragraph()
    rel_para.paragraph_format.space_before = Pt(8)
    rel_para.paragraph_format.space_after = Pt(2)
    r_title = rel_para.add_run(title)
    r_title.bold = True
    r_title.font.color.rgb = NAVY
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(10.5)

    exp_para = doc.add_paragraph()
    exp_para.paragraph_format.left_indent = Inches(0.3)
    exp_para.paragraph_format.space_after = Pt(2)
    er = exp_para.add_run(explanation)
    er.font.name = "Calibri"
    er.font.size = Pt(10)
    er.font.color.rgb = DARK

    key_para = doc.add_paragraph()
    key_para.paragraph_format.left_indent = Inches(0.3)
    key_para.paragraph_format.space_after = Pt(4)
    kr1 = key_para.add_run("\u25b6  Join key: ")
    kr1.bold = True
    kr1.font.color.rgb = TEAL
    kr1.font.name = "Calibri"
    kr1.font.size = Pt(9.5)
    kr2 = key_para.add_run(join_key)
    kr2.italic = True
    kr2.font.name = "Calibri"
    kr2.font.size = Pt(9.5)
    kr2.font.color.rgb = GREY

# ── 7.5 Limitations ──────────────────────────────────────────────────────────
section_heading(doc, "7.5  Data Limitations to Know for the Interview")

limitations = [
    ("No unique record IDs",
     "Tables share conceptual dimensions (year, income band) but have no formal primary keys "
     "or foreign keys. Joins are done by matching labels — which means label format must match exactly "
     "(e.g. '<$35K' vs '$35K\u2013$50K' are different levels)."),
    ("Income bands are not consistent across files",
     "Demographics uses '$35K\u2013$50K'. Rental uses '<$50K' (cumulative threshold). "
     "These overlap but are not identical. The dashboard handles this, but it is important "
     "to note when comparing numbers across files."),
    ("Years covered differ by table",
     "Some tables run 2013\u20132024, others 2015\u20132024, others show only 3 snapshot years "
     "(2014/2019/2024). You cannot always build a continuous trend from every variable."),
    ("No sub-county geography",
     "All data is at the Wake County level. There is no census tract or zip code breakdown "
     "in these files \u2014 meaning the data cannot show which neighbourhoods are worst affected."),
    ("Cumulative vs. band income thresholds",
     "The rental deficit figures use cumulative thresholds (<$35K means all households earning "
     "up to $35K). Demographics income bands are ranges ($20K\u2013$35K). "
     "The numbers are not directly comparable without conversion."),
]

lim_tbl = doc.add_table(rows=1, cols=2)
lim_tbl.style = "Table Grid"
for cell, h in zip(lim_tbl.rows[0].cells, ["Limitation", "What It Means in Practice"]):
    set_cell_bg(cell, "1B3A6B")
    set_cell_borders(cell, "FFFFFF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)

for idx, (lim, detail) in enumerate(limitations):
    row_l = lim_tbl.add_row()
    bg_l = "FFF8E8" if idx % 2 == 0 else "FFFDF5"
    for ci, (cell, val) in enumerate(zip(row_l.cells, [lim, detail])):
        set_cell_bg(cell, bg_l)
        set_cell_borders(cell, "E0C060")
        cell.width = Inches(1.8) if ci == 0 else Inches(5.5)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(val)
        r.bold = (ci == 0)
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = NAVY if ci == 0 else DARK

doc.add_paragraph()

# Final note
note_para = doc.add_paragraph()
note_para.paragraph_format.space_before = Pt(8)
shade_para(note_para, "1B3A6B")
note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = note_para.add_run(
    "Section 7 added April 2026  \u2014  "
    "Data model covers all 26 tables across 4 wakehousingdata.org Excel files"
)
nr.font.size = Pt(8)
nr.font.color.rgb = WHITE
nr.font.name = "Calibri"
nr.italic = True

doc.save(PATH)
print(f"\u2713 Section 7 appended: {PATH}")
